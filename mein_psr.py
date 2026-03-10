import tkinter as tk
from tkinter import messagebox
from pynput import mouse
import mss
import os
import time
import ctypes
from datetime import datetime
from PIL import Image, ImageDraw
from docx import Document
from docx.shared import Inches

# 1. Hauptordner vorbereiten
haupt_ordner = "Meine_Aufzeichnung"
os.makedirs(haupt_ordner, exist_ok=True)

# Globale Variablen
aktueller_ordner = ""
docx_pfad = ""
dokument = None
schritt = 1
letzter_klick_zeit = 0

# Steuerungs-Variablen
session_aktiv = False
aufzeichnung_laeuft = False
lupe_aktiv = True
aktuelle_sprache = "DE" # Startsprache

# === DAS WÖRTERBUCH (Alle Texte in DE und EN) ===
TEXTE = {
    "DE": {
        "btn_start_neu": "▶️ Neue Aufzeichnung",
        "btn_pause": "⏸ Pausieren",
        "btn_weiter": "▶️ Weiter",
        "btn_speichern": "💾 Anleitung abschließen",
        "chk_lupe": "Detail-Lupe aktivieren",
        "btn_hilfe": "❓ Hilfe & Anleitung",
        "btn_beenden": "❌ Tool komplett beenden",
        "status_bereit": "Status: ⏹ Bereit für Neues",
        "status_laeuft": "Status: 🔴 Aufzeichnung läuft...",
        "status_pausiert": "Status: ⏸ Pausiert",
        "status_gespeichert": "Status: 💾 Gespeichert & Bereit",
        "doc_titel": "Meine Schrittaufzeichnung",
        "doc_erstellt": "Erstellt am:",
        "doc_autor": "Autor:",
        "doc_tipp": "Tipp: Unter jedem Bild kannst du eigene Notizen ergänzen.",
        "doc_schritt": "Schritt",
        "doc_programm": "Aktives Programm:",
        "doc_lupe": "Detailansicht (Lupe):",
        "doc_notiz": "[Hier klicken für eigene Notizen...]",
        "klick_links": "Linker Mausklick",
        "klick_rechts": "Rechter Mausklick",
        "klick_mitte": "Mittlere Maustaste / Anderer Klick",
        "klick_doppel": "Doppelklick",
        "fenster_unbekannt": "Unbekanntes Fenster",
        "hilfe_titel": "Hilfe & Anleitung",
        "hilfe_text": "Willkommen zum PSR Tool!\n\n▶️ Neue Aufzeichnung: Startet eine neue Session.\n⏸ Pausieren: Pausiert die Aufnahme für private Klicks.\n🔍 Detail-Lupe: Fügt einen vergrößerten Ausschnitt deines Klicks hinzu.\n💾 Anleitung abschließen: Speichert das fertige Word-Dokument.\n\nEntwickelt von Sacha Schwendimann.",
        "btn_lang": "🇬🇧 Switch to English"
    },
    "EN": {
        "btn_start_neu": "▶️ New Recording",
        "btn_pause": "⏸ Pause",
        "btn_weiter": "▶️ Continue",
        "btn_speichern": "💾 Finish Recording",
        "chk_lupe": "Enable Detail Magnifier",
        "btn_hilfe": "❓ Help & Instructions",
        "btn_beenden": "❌ Exit Tool",
        "status_bereit": "Status: ⏹ Ready for new recording",
        "status_laeuft": "Status: 🔴 Recording in progress...",
        "status_pausiert": "Status: ⏸ Paused",
        "status_gespeichert": "Status: 💾 Saved & Ready",
        "doc_titel": "My Step Recording",
        "doc_erstellt": "Created at:",
        "doc_autor": "Author:",
        "doc_tipp": "Tip: You can add your own notes under each image.",
        "doc_schritt": "Step",
        "doc_programm": "Active Program:",
        "doc_lupe": "Detail view (Magnifier):",
        "doc_notiz": "[Click here to add your own notes...]",
        "klick_links": "Left Click",
        "klick_rechts": "Right Click",
        "klick_mitte": "Middle Button / Other Click",
        "klick_doppel": "Double Click",
        "fenster_unbekannt": "Unknown Window",
        "hilfe_titel": "Help & Instructions",
        "hilfe_text": "Welcome to the PSR Tool!\n\n▶️ New Recording: Starts a new session.\n⏸ Pause: Pauses the recording for private clicks.\n🔍 Detail Magnifier: Adds a zoomed-in cutout of your click.\n💾 Finish Recording: Saves the final Word document.\n\nDeveloped by Sacha Schwendimann.",
        "btn_lang": "🇩🇪 Auf Deutsch wechseln"
    }
}

def t(key):
    return TEXTE[aktuelle_sprache][key]

# === FEATURE 1: Fenster-Spion ===
def get_aktives_fenster():
    try:
        hwnd = ctypes.windll.user32.GetForegroundWindow()
        length = ctypes.windll.user32.GetWindowTextLengthW(hwnd)
        buf = ctypes.create_unicode_buffer(length + 1)
        ctypes.windll.user32.GetWindowTextW(hwnd, buf, length + 1)
        return buf.value if buf.value else t("fenster_unbekannt")
    except:
        return t("fenster_unbekannt")

# Maus-Zuhörer
def bei_klick(x, y, button, pressed):
    global schritt, letzter_klick_zeit
    
    if not (session_aktiv and aufzeichnung_laeuft):
        return

    if pressed:
        aktuelle_zeit = time.time()
        zeit_differenz = aktuelle_zeit - letzter_klick_zeit
        kreis_farbe = "blue" 
        
        if button == mouse.Button.left:
            klick_art = t("klick_links")
            kreis_farbe = "red"
        elif button == mouse.Button.right:
            klick_art = t("klick_rechts")
            kreis_farbe = "orange"
        else:
            klick_art = t("klick_mitte")

        if zeit_differenz < 0.4:
            klick_art = f'{t("klick_doppel")} ({klick_art})'

        letzter_klick_zeit = aktuelle_zeit
        fenster_titel = get_aktives_fenster()
        
        with mss.mss() as sct:
            monitor_nummer = 1 
            monitor_daten = sct.monitors[1] 
            
            for i in range(1, len(sct.monitors)):
                m = sct.monitors[i]
                if m["left"] <= x < m["left"] + m["width"] and m["top"] <= y < m["top"] + m["height"]:
                    monitor_nummer = i
                    monitor_daten = m
                    break
            
            bild_name = f"schritt_{schritt}.png"
            bild_pfad = os.path.join(aktueller_ordner, bild_name)
            
            sct_img = sct.grab(monitor_daten)
            img = Image.frombytes("RGB", sct_img.size, sct_img.bgra, "raw", "BGRX")
            zeichnen = ImageDraw.Draw(img)
            
            punkt_x = x - monitor_daten["left"]
            punkt_y = y - monitor_daten["top"]
            
            radius = 25
            zeichnen.ellipse((punkt_x - radius, punkt_y - radius, punkt_x + radius, punkt_y + radius), outline=kreis_farbe, width=5)
            cursor_punkte = [(punkt_x, punkt_y), (punkt_x, punkt_y + 17), (punkt_x + 5, punkt_y + 13), (punkt_x + 11, punkt_y + 11)]
            zeichnen.polygon(cursor_punkte, fill="white", outline="black")
            
            img.save(bild_pfad)
            
            dokument.add_heading(f'{t("doc_schritt")} {schritt}: {klick_art}', level=2)
            dokument.add_paragraph(f'{t("doc_programm")} {fenster_titel}')
            dokument.add_picture(bild_pfad, width=Inches(6.0))
            
            if lupe_aktiv:
                lupe_groesse = 150 
                l_links = max(0, punkt_x - lupe_groesse)
                l_oben = max(0, punkt_y - lupe_groesse)
                l_rechts = min(img.width, punkt_x + lupe_groesse)
                l_unten = min(img.height, punkt_y + lupe_groesse)
                
                lupe_img = img.crop((l_links, l_oben, l_rechts, l_unten))
                lupe_pfad = os.path.join(aktueller_ordner, f"schritt_{schritt}_lupe.png")
                lupe_img.save(lupe_pfad)
                
                dokument.add_paragraph(t("doc_lupe"))
                dokument.add_picture(lupe_pfad, width=Inches(2.5))
            
            dokument.add_paragraph(t("doc_notiz"))
            dokument.add_paragraph('') 
            dokument.save(docx_pfad)
            schritt += 1

# === DAS MENÜ (GUI) ===

def update_ui_texte():
    btn_sprache.config(text=t("btn_lang"))
    chk_lupe.config(text=t("chk_lupe"))
    btn_speichern.config(text=t("btn_speichern"))
    btn_hilfe.config(text=t("btn_hilfe"))
    btn_beenden.config(text=t("btn_beenden"))
    
    if not session_aktiv:
        btn_start_pause.config(text=t("btn_start_neu"))
        if "Gespeichert" in lbl_status.cget("text") or "Saved" in lbl_status.cget("text"):
            lbl_status.config(text=t("status_gespeichert"))
        else:
            lbl_status.config(text=t("status_bereit"))
    else:
        if aufzeichnung_laeuft:
            btn_start_pause.config(text=t("btn_pause"))
            lbl_status.config(text=t("status_laeuft"))
        else:
            btn_start_pause.config(text=t("btn_weiter"))
            lbl_status.config(text=t("status_pausiert"))

def sprache_umschalten():
    global aktuelle_sprache
    if aktuelle_sprache == "DE":
        aktuelle_sprache = "EN"
    else:
        aktuelle_sprache = "DE"
    update_ui_texte()

def start_pause_klick():
    global session_aktiv, aufzeichnung_laeuft, aktueller_ordner, docx_pfad, dokument, schritt
    
    if not session_aktiv:
        session_aktiv = True
        aufzeichnung_laeuft = True
        schritt = 1 
        
        zeitstempel = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        aktueller_ordner = os.path.join(haupt_ordner, zeitstempel)
        os.makedirs(aktueller_ordner, exist_ok=True)
        
        docx_pfad = os.path.join(aktueller_ordner, "protokoll.docx")
        dokument = Document()
        dokument.add_heading(t("doc_titel"), 0)
        dokument.add_paragraph(f'{t("doc_erstellt")} {datetime.now().strftime("%d.%m.%Y - %H:%M:%S")}')
        dokument.add_paragraph(f'{t("doc_autor")} Sacha Schwendimann')
        dokument.add_paragraph(t("doc_tipp"))
        dokument.save(docx_pfad)
        
        btn_start_pause.config(text=t("btn_pause"), bg="#ffcc00")
        btn_speichern.config(state="normal", bg="#66ccff") 
        lbl_status.config(text=t("status_laeuft"), fg="red")

    elif aufzeichnung_laeuft:
        aufzeichnung_laeuft = False
        btn_start_pause.config(text=t("btn_weiter"), bg="#99ff99")
        lbl_status.config(text=t("status_pausiert"), fg="orange")
        
    else:
        aufzeichnung_laeuft = True
        btn_start_pause.config(text=t("btn_pause"), bg="#ffcc00")
        lbl_status.config(text=t("status_laeuft"), fg="red")

def speichern_und_abschliessen():
    global session_aktiv, aufzeichnung_laeuft
    
    if session_aktiv:
        session_aktiv = False
        aufzeichnung_laeuft = False
        
        btn_start_pause.config(text=t("btn_start_neu"), bg="#99ff99")
        btn_speichern.config(state="disabled", bg="#cccccc") 
        lbl_status.config(text=t("status_gespeichert"), fg="green")

def lupe_umschalten():
    global lupe_aktiv
    lupe_aktiv = check_var_lupe.get()

def zeige_hilfe():
    messagebox.showinfo(t("hilfe_titel"), t("hilfe_text"))

def programm_beenden():
    try:
        maus_zuhörer.stop()
    except:
        pass
    fenster.destroy()

# Fenster erstellen
fenster = tk.Tk()
fenster.title("PSR Tool - Sacha Schwendimann")
fenster.geometry("280x350") 
fenster.attributes('-topmost', True)
fenster.configure(padx=10, pady=10)

# Sprachen-Button 
btn_sprache = tk.Button(fenster, text="🇬🇧 Switch to English", bg="#e0e0e0", font=("Arial", 8), command=sprache_umschalten)
btn_sprache.pack(anchor="ne", pady=(0, 5)) 

# Status-Text
lbl_status = tk.Label(fenster, text="Status: ⏹ Bereit für Neues", font=("Arial", 10, "bold"))
lbl_status.pack(pady=5)

# Start/Pause Button
btn_start_pause = tk.Button(fenster, text="▶️ Neue Aufzeichnung", bg="#99ff99", font=("Arial", 10, "bold"), command=start_pause_klick, width=22, height=2)
btn_start_pause.pack(pady=5)

# Speichern & Abschließen Button
btn_speichern = tk.Button(fenster, text="💾 Anleitung abschließen", bg="#cccccc", command=speichern_und_abschliessen, width=22, state="disabled")
btn_speichern.pack(pady=5)

# Checkbox für die Lupe
check_var_lupe = tk.BooleanVar(value=True)
chk_lupe = tk.Checkbutton(fenster, text="Detail-Lupe aktivieren", variable=check_var_lupe, command=lupe_umschalten)
chk_lupe.pack(pady=5)

# Hilfe-Button
btn_hilfe = tk.Button(fenster, text="❓ Hilfe & Anleitung", bg="#e6e6e6", command=zeige_hilfe, width=22)
btn_hilfe.pack(pady=5)

# Programm Beenden Button 
btn_beenden = tk.Button(fenster, text="❌ Tool komplett beenden", bg="#ff9999", command=programm_beenden, width=22)
btn_beenden.pack(pady=5)

# Persönliche Signatur
lbl_autor = tk.Label(fenster, text="© Sacha Schwendimann", font=("Arial", 8, "italic"), fg="gray")
lbl_autor.pack(side="bottom", pady=5)

# Initial einmal die Texte setzen
update_ui_texte()

# Maus-Zuhörer im Hintergrund starten
maus_zuhörer = mouse.Listener(on_click=bei_klick)
maus_zuhörer.start()

# Das Menü anzeigen
fenster.mainloop()